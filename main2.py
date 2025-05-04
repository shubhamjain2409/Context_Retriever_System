import os
from langchain_community.document_loaders import Docx2txtLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.llms import HuggingFaceHub
from langchain.chains import RetrievalQA
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from langchain.schema import Document

from docx import Document as DocxDocument
from langchain.schema import Document
import re

def read_docx_to_text(docx_path):
    doc = DocxDocument(docx_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return full_text

# Simple regex-based sentence tokenizer
def regex_sent_tokenize(text):
    return re.split(r'(?<=[.!?])\s+', text)

def split_by_section_and_sentences(text, sentences_per_chunk=5):
    sections = re.split(r'(?=\n\d+\.\s)', text)
    chunks = []

    for section in sections:
        lines = section.strip().split('\n', 1)
        title = lines[0].strip()
        content = lines[1].strip() if len(lines) > 1 else ""

        sentences = regex_sent_tokenize(content)
        for i in range(0, len(sentences), sentences_per_chunk):
            chunk_text = " ".join(sentences[i:i + sentences_per_chunk])
            chunks.append(Document(page_content=f"{title}\n{chunk_text}", metadata={"section": title}))

    return chunks

# === USAGE ===
docx_file_path = "/content/drive/My Drive/LLMda_Data/Applied_AI_Engineer/Sample_Sources/USB3_IP_User_Guide_v2.docx"
text = read_docx_to_text(docx_file_path)
chunks = split_by_section_and_sentences(text, sentences_per_chunk=5)

class TfidfRetriever:
  def __init__(self, documents):
      self.texts = [doc.page_content for doc in documents]
      self.docs = documents
      self.vectorizer = TfidfVectorizer()
      self.tfidf_matrix = self.vectorizer.fit_transform(self.texts)

  def get_relevant_documents(self, query, ip_core_name):
      query = query.replace(ip_core_name, '')
      print(query)
      query_vec = self.vectorizer.transform([query])
      cosine_scores = cosine_similarity(query_vec, self.tfidf_matrix).flatten()
      top_indices = cosine_scores.argsort()[-3:][::-1]
      return [self.docs[i] for i in top_indices]


sparse = TfidfRetriever(chunks)

results = sparse.get_relevant_documents("What features does the Ethernet 200G IP Core provide for error detection and correction?","Ethernet 200G IP Core")

paragraphs = [doc.page_content.strip() for doc in results]

for i, para in enumerate(paragraphs, 1):
    print(f"Paragraph {i}:\n{para}\n")