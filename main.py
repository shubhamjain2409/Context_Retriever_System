import os
from langchain.document_loaders import TextLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document

from docx import Document as DocxDocument
from langchain.schema import Document
import re

def read_docx_to_text(docx_path):
    doc = DocxDocument(docx_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return full_text

# Simple regex-based sentence tokenizer
def regex_sent_tokenize(text):
    return re.split(r'(?<=[.!?])\s+', text)

def split_by_section_and_sentences(text, sentences_per_chunk=5):
    sections = re.split(r'(?=\n\d+\.\s)', text)
    chunks = []

    for section in sections:
        lines = section.strip().split('\n', 1)
        title = lines[0].strip()
        content = lines[1].strip() if len(lines) > 1 else ""

        sentences = regex_sent_tokenize(content)
        for i in range(0, len(sentences), sentences_per_chunk):
            chunk_text = " ".join(sentences[i:i + sentences_per_chunk])
            chunks.append(Document(page_content=f"{title}\n{chunk_text}", metadata={"section": title}))

    return chunks

# === USAGE ===
docx_file_path = "USB3_IP_User_Guide_v2.docx"
text = read_docx_to_text(docx_file_path)
chunks = split_by_section_and_sentences(text, sentences_per_chunk=5)

def create_embedding_retriever(chunks):
    from langchain.vectorstores import Chroma
    from langchain.embeddings import HuggingFaceEmbeddings
    from langchain.schema import Document
    import uuid

    def query_preprocess(query: str) -> str:
        return f"query: {query}"

    def passage_preprocess(doc: Document) -> Document:
        doc.page_content = f"passage: {doc.page_content}"
        return doc

    chunks = [passage_preprocess(chunk) for chunk in chunks]

    embeddings = HuggingFaceEmbeddings(
        model_name="intfloat/e5-large-v2",
        encode_kwargs={'normalize_embeddings': True}
    )

    # Create a unique, fresh Chroma collection (no persist_directory = in-memory)
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=f"e5_collection_{uuid.uuid4()}"  # ensures a fresh collection
    )

    retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 3})

    class E5Retriever:
        def __init__(self, retriever):
            self.retriever = retriever

        def get_relevant_documents(self, query):
            return self.retriever.get_relevant_documents(query_preprocess(query))

    return E5Retriever(retriever)


retriever = create_embedding_retriever(chunks)

def get_top_k_relevant_paragraphs(query, retriever, k):
    
    return retriever.get_relevant_documents(query)

query = "How does the USB 3.0 IP core handle link error recovery?"
top_docs = get_top_k_relevant_paragraphs(query, retriever, "PCIe Gen 5 x8 IP core", 3)


paragraphs = [doc.page_content.strip() for doc in top_docs]

for i, para in enumerate(paragraphs, 1):
    print(f"Paragraph {i}:\n{para}\n")


if __name__ == "__main__":
    main()